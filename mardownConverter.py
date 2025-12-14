#!/usr/bin/env python
import re
import json
from io import BytesIO
from pathlib import Path
from textwrap import dedent

from docx import Document
import mammoth
from bs4 import BeautifulSoup
from html_to_markdown import convert_to_markdown

# --------------------------------------------------
# Flask setup
# --------------------------------------------------


FORM_TEMPLATE = """
<!doctype html>
<html lang="en">
  <head>
    <meta charset="utf-8">
    <title>UPN DOCX → Markdown Converter</title>
    <meta name="viewport" content="width=device-width, initial-scale=1">
    <style>
      :root {
        --bg: #0f172a;
        --card-bg: #0b1220;
        --accent: #38bdf8;
        --accent-soft: rgba(56, 189, 248, 0.12);
        --border-subtle: #1e293b;
        --text-main: #e5e7eb;
        --text-muted: #9ca3af;
        --error: #f97373;
        --radius-lg: 14px;
      }

      * {
        box-sizing: border-box;
      }

      body {
        margin: 0;
        min-height: 100vh;
        font-family: system-ui, -apple-system, BlinkMacSystemFont, "Segoe UI",
          sans-serif;
        color: var(--text-main);
        background:
          radial-gradient(circle at top left, rgba(56,189,248,0.25), transparent 55%),
          radial-gradient(circle at bottom right, rgba(129,140,248,0.22), transparent 55%),
          var(--bg);
        display: flex;
        align-items: center;
        justify-content: center;
        padding: 2.5rem 1.5rem;
      }

      .card {
        width: 100%;
        max-width: 520px;
        background: linear-gradient(145deg, #020617, var(--card-bg));
        border-radius: var(--radius-lg);
        border: 1px solid var(--border-subtle);
        box-shadow:
          0 24px 60px rgba(15,23,42,0.75),
          0 0 0 1px rgba(148,163,184,0.05);
        padding: 1.8rem 1.8rem 1.9rem;
      }

      .card-header {
        display: flex;
        align-items: center;
        gap: 0.8rem;
        margin-bottom: 1.2rem;
      }

      .card-pill {
        width: 32px;
        height: 32px;
        border-radius: 999px;
        background: radial-gradient(circle at 30% 20%, #facc15, #f97316);
        display: flex;
        align-items: center;
        justify-content: center;
        box-shadow: 0 0 0 1px rgba(248,250,252,0.25);
        font-size: 0.85rem;
        font-weight: 600;
        color: #020617;
      }

      h1 {
        font-size: 1.3rem;
        letter-spacing: 0.02em;
        margin: 0;
      }

      .subtitle {
        margin-top: 0.15rem;
        font-size: 0.8rem;
        color: var(--text-muted);
      }

      form {
        margin-top: 1rem;
      }

      .field {
        margin-top: 1rem;
      }

      .field-label {
        display: flex;
        justify-content: space-between;
        align-items: center;
        margin-bottom: 0.25rem;
        font-size: 0.82rem;
        color: var(--text-muted);
      }

      .field-label span {
        font-weight: 500;
        color: var(--text-main);
      }

      .hint {
        font-size: 0.75rem;
        color: var(--text-muted);
      }

      input[type="text"],
      input[type="file"] {
        width: 100%;
        font-size: 0.9rem;
        padding: 0.55rem 0.7rem;
        border-radius: 0.6rem;
        border: 1px solid var(--border-subtle);
        background-color: rgba(15,23,42,0.72);
        color: var(--text-main);
        outline: none;
        transition: border-color 0.16s ease, box-shadow 0.16s ease, background-color 0.16s ease;
      }

      input[type="file"] {
        padding: 0.45rem 0.7rem;
      }

      input::file-selector-button {
        margin-right: 0.7rem;
        border: none;
        border-radius: 999px;
        padding: 0.4rem 0.8rem;
        background-color: var(--accent-soft);
        color: var(--accent);
        font-size: 0.78rem;
        cursor: pointer;
      }

      input[type="text"]:focus,
      input[type="file"]:focus-visible {
        border-color: var(--accent);
        box-shadow: 0 0 0 1px rgba(56,189,248,0.55);
        background-color: rgba(15,23,42,0.96);
      }

      button[type="submit"] {
        margin-top: 1.5rem;
        width: 100%;
        border: none;
        border-radius: 999px;
        padding: 0.7rem 1rem;
        font-size: 0.95rem;
        font-weight: 500;
        letter-spacing: 0.01em;
        cursor: pointer;
        background: linear-gradient(135deg, #0ea5e9, #6366f1);
        color: #f9fafb;
        box-shadow:
          0 12px 30px rgba(37,99,235,0.45),
          0 0 0 1px rgba(15,23,42,0.6);
        transition: transform 0.12s ease, box-shadow 0.12s ease, filter 0.12s ease;
      }

      button[type="submit"]:hover {
        filter: brightness(1.03);
        transform: translateY(-1px);
        box-shadow:
          0 18px 40px rgba(37,99,235,0.6),
          0 0 0 1px rgba(15,23,42,0.7);
      }

      button[type="submit"]:active {
        transform: translateY(0);
        box-shadow:
          0 12px 26px rgba(30,64,175,0.7),
          0 0 0 1px rgba(15,23,42,0.7);
      }

      .error {
        margin-top: 1rem;
        padding: 0.6rem 0.75rem;
        border-radius: 0.6rem;
        background-color: rgba(239,68,68,0.08);
        border: 1px solid rgba(248,113,113,0.4);
        color: var(--error);
        font-size: 0.82rem;
      }

      .footer-note {
        margin-top: 1.1rem;
        font-size: 0.74rem;
        color: var(--text-muted);
        text-align: center;
      }

      .footer-note code {
        font-size: 0.72rem;
        background-color: rgba(15,23,42,0.85);
        border-radius: 0.35rem;
        padding: 0.1rem 0.35rem;
      }
    </style>
  </head>
  <body>
    <main class="card">
      <div class="card-header">
        <div class="card-pill">UPN</div>
        <div>
          <h1>DOCX → Markdown Converter</h1>
          <p class="subtitle">Generate Academy-ready Markdown + images from a Word UPN template.</p>
        </div>
      </div>

      <form method="post" enctype="multipart/form-data">
        <div class="field">
          <div class="field-label">
            <span>Work Item Number</span>
            <span class="hint">e.g. <code>AWI00123456</code></span>
          </div>
          <input type="text" name="work_item_number" required>
        </div>

        <div class="field">
          <div class="field-label">
            <span>Language</span>
            <span class="hint">Use 2-letter code (e.g. <code>EN</code>)</span>
          </div>
          <input type="text" name="language" maxlength="5" required>
        </div>

        <div class="field">
          <div class="field-label">
            <span>DOCX file</span>
            <span class="hint">Upload the UPN Word template</span>
          </div>
          <input type="file" name="docx_file" accept=".docx" required>
        </div>

        <button type="submit">Convert and Download ZIP</button>
      </form>

      {% if error %}
        <p class="error">{{ error }}</p>
      {% endif %}

      <p class="footer-note">
        Output: <code>{{ '{{' }}work_item_number{{ '}}' }}-{{ '{{' }}language{{ '}}' }}-{{ '{{' }}language{{ '}}' }}.md</code> + images in <code>_images/</code>.
      </p>
    </main>
  </body>
</html>
"""


# --------------------------------------------------
# Utility functions (metadata, YAML front matter)
# --------------------------------------------------

def parse_bool(text: str, default: bool = False) -> bool:
    if text is None:
        return default
    t = text.strip().lower()
    if t in ("true", "yes", "y", "1"):
        return True
    if t in ("false", "no", "n", "0"):
        return False
    return default


def parse_int(text: str, default: int = 0) -> int:
    if text is None or not text.strip():
        return default
    try:
        return int(text.strip())
    except ValueError:
        return default


def parse_search_terms(text: str):
    """
    Convert e.g. "20250723e, some, search, terms"
    into ['20250723e', 'some', 'search', 'terms']
    """
    if text is None:
        return []

    raw = text.strip()

    # If it looks like a Python list ['a','b'], strip brackets
    if raw.startswith("[") and raw.endswith("]"):
        raw = raw[1:-1]

    parts = [p.strip() for p in raw.split(",")]
    return [p.strip(" '\"") for p in parts if p]


def extract_metadata(doc: Document) -> dict:
    """
    Reads the FIRST table and builds a metadata dict.
    """
    metadata = {}
    if not doc.tables:
        return metadata

    table = doc.tables[0]
    for row in table.rows:
        key = row.cells[0].text.strip()
        value = row.cells[1].text.strip()
        if key:
            metadata[key] = value

    yaml_meta = {}
    version_numbers = {}
    version_keys = {"software_version", "dpr", "std", "gp1", "gp2"}

    for key, value in metadata.items():
        if key in version_keys:
            version_numbers[key] = value
        else:
            yaml_meta[key] = value

    yaml_meta["version_numbers"] = version_numbers

    # ---- Apply conversions ----
    yaml_meta["published"] = parse_bool(yaml_meta.get("published", "true"), default=True)
    yaml_meta["access_warning"] = parse_bool(yaml_meta.get("access_warning", "false"), default=False)
    yaml_meta["collect_feedback"] = parse_bool(yaml_meta.get("collect_feedback", "true"), default=True)

    yaml_meta["duration"] = parse_int(yaml_meta.get("duration", "7"), default=7)
    yaml_meta["search_terms"] = parse_search_terms(yaml_meta.get("search_terms", ""))

    # Defaults
    yaml_meta.setdefault("product", "CargoWise")
    yaml_meta.setdefault("course_series", "Product learning")
    yaml_meta.setdefault("membership_level", "Software User")
    yaml_meta.setdefault("summary", "")

    return yaml_meta


def build_front_matter(meta: dict) -> str:
    header_comment = dedent("""\
    ###########################################################################################################
    # This is the official update note template, all future templates should be based on this as a reference. #
    ###########################################################################################################
    """)

    terms = meta.get("search_terms", [])
    search_terms_yaml = "[" + ", ".join(f"'{t}'" for t in terms) + "]"

    summary_text = meta.get("summary", "").rstrip()
    summary_lines = summary_text.splitlines() if summary_text else []
    indented_summary = "\n".join(
        ("  " + line) if line else ""
        for line in summary_lines
    )
    if not indented_summary:
        indented_summary = "  "

    versions = meta.get("version_numbers", {})

    lines = []
    lines.append("---")
    lines.append(header_comment.rstrip())
    lines.append("")
    lines.append("# Required front matter in YAML")
    lines.append(f"published: {str(meta.get('published', True)).lower()}")
    lines.append(f"title: {meta.get('title', 'Document Title')}")
    lines.append(f"product: {meta.get('product', 'CargoWise')}")
    lines.append(f"search_terms: {search_terms_yaml}")
    lines.append("summary: |")
    lines.append(indented_summary)
    lines.append(f"access_warning: {str(meta.get('access_warning', False)).lower()}")
    lines.append(f"release_date: {meta.get('release_date', '2025-11-12')}")
    lines.append(f"reissue_date: \"{meta.get('reissue_date', '')}\"")
    lines.append(f"published_date: {meta.get('published_date', '2025-11-12')}")
    lines.append(f"duration: {meta.get('duration', 7)}")
    lines.append(f"course_series: {meta.get('course_series', 'Product learning')}")
    lines.append(f"membership_level: {meta.get('membership_level', 'Software User')}")
    lines.append("")
    lines.append("version_numbers:")
    lines.append(f"  software_version: \"{versions.get('software_version', '')}\"")
    lines.append(f"  dpr: \"{versions.get('dpr', '')}\"")
    lines.append(f"  std: \"{versions.get('std', '')}\"")
    lines.append(f"  gp1: \"{versions.get('gp1', '')}\"")
    lines.append(f"  gp2: \"{versions.get('gp2', '')}\"")
    lines.append("")
    lines.append(f"collect_feedback: {str(meta.get('collect_feedback', True)).lower()}")
    lines.append("")
    lines.append('editing_division: ""')
    lines.append('editing_department: ""')
    lines.append('study_access: ""')
    lines.append('study_division: ""')
    lines.append('study_department: ""')
    lines.append("---")

    return "\n".join(lines) + "\n\n"

# --------------------------------------------------
# Image converter (in-memory)
# --------------------------------------------------

def make_image_converter(image_store: dict, wi: str, lang: str):
    """
    image_store: dict[path_in_zip -> bytes]

    Paths will look like:
      _images/{WI}-{LANG}-{LANG}-001.png
    """
    from itertools import count
    counter = count(1)

    def convert_image(image):
        image_number = next(counter)

        ext_map = {
            "image/png": ".png",
            "image/jpeg": ".jpg",
            "image/jpg": ".jpg",
            "image/gif": ".gif",
            "image/tiff": ".tiff",
            "image/bmp": ".bmp",
        }
        ext = ext_map.get(image.content_type, ".bin")

        rel_path = f"_images/{wi}-{lang}-{lang}-{image_number:03d}{ext}"

        with image.open() as image_bytes:
            data = image_bytes.read()

        image_store[rel_path] = data

        return {"src": rel_path}

    return mammoth.images.img_element(convert_image)

# --------------------------------------------------
# Markdown post-processing helpers
# --------------------------------------------------

def fix_broken_images(md: str) -> str:
    pattern = re.compile(
        r'!\[(.*?)\n+([^\]]*?)\]\(([^)]+)\)',
        re.DOTALL
    )

    def repl(m):
        alt1 = m.group(1).strip()
        alt2 = m.group(2).strip()
        path = m.group(3).strip()
        alt = " ".join([p for p in [alt1, alt2] if p])
        return f'![{alt}]({path})'

    return pattern.sub(repl, md)


def fix_toc_and_anchors(md: str) -> str:
    lines = md.splitlines()
    toc_ids = []

    toc_link_re = re.compile(
        r'\[(?P<label>[^\]]*?)(?:\s+\d+)?\]\(#(?P<id>_Toc\d+)\)'
    )

    cleaned_lines = []
    for line in lines:
        def repl(m):
            label = m.group('label').strip()
            anchor_id = m.group('id')
            toc_ids.append(anchor_id)
            return f'[{label}](#{anchor_id})'

        new_line = toc_link_re.sub(repl, line)
        cleaned_lines.append(new_line)

    lines = cleaned_lines

    heading_indices = []
    for i in range(len(lines) - 1):
        title = lines[i].strip()
        underline = lines[i + 1].strip()
        if title and re.fullmatch(r'-{3,}|={3,}', underline):
            heading_indices.append(i)

    anchor_by_line = {}
    for idx, heading_idx in enumerate(heading_indices):
        if idx < len(toc_ids):
            anchor_by_line[heading_idx] = toc_ids[idx]

    final_lines = []
    for i, line in enumerate(lines):
        if i in anchor_by_line:
            final_lines.append(f'<a id="{anchor_by_line[i]}"></a>')
        final_lines.append(line)

    return "\n".join(final_lines)


def _make_callout_block(kind: str, inner_md: str) -> str:
    lines = [ln.rstrip() for ln in inner_md.splitlines()]

    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()

    if not lines:
        return ""

    out = [f"> [!{kind}]"]
    for ln in lines:
        if ln.strip():
            out.append("> " + ln)
        else:
            out.append(">")

    return "\n".join(out) + "\n\n"


def apply_callout_divs(html: str) -> str:
    soup = BeautifulSoup(html, "html.parser")

    def process_div(class_name: str, kind: str) -> None:
        for div in soup.find_all("div", class_=class_name):
            inner_html = "".join(str(child) for child in div.children)
            inner_md = convert_to_markdown(inner_html)
            callout_md = _make_callout_block(kind, inner_md)
            div.replace_with(callout_md)

    process_div("upn-note", "NOTE")
    process_div("upn-important", "IMPORTANT")
    process_div("upn-action", "ACTION")

    return str(soup)


def fix_escaped_callouts(md: str) -> str:
    lines = md.splitlines()
    out = []

    for line in lines:
        stripped = line.lstrip()

        if stripped.startswith(r'\> \[!NOTE\]'):
            out.append('> [!NOTE]')
        elif stripped.startswith(r'\> \[!IMPORTANT\]'):
            out.append('> [!IMPORTANT]')
        elif stripped.startswith(r'\> \[!ACTION\]'):
            out.append('> [!ACTION]')
        elif stripped.startswith(r'\> '):
            prefix_len = len(line) - len(stripped)
            out.append(' ' * prefix_len + '> ' + stripped[3:])
        else:
            out.append(line)

    return "\n".join(out)


def pseudo_pretty_json(raw: str) -> str:
    s = re.sub(r',\s{4,}', ',\n  ', raw)
    s = re.sub(r'\{\s{4,}', '{\n  ', s)
    s = re.sub(r'\s{4,}\}', '\n}', s)
    s = re.sub(r'(?<!\n)[ \t]{2,}', ' ', s)
    return s


def fix_code_blocks(md: str) -> str:
    lines = md.splitlines()
    out = []

    heading_to_lang = {
        "JSON Code": "json",
        "JSON": "json",
        "Python Code": "python",
        "Python": "python",
        "PowerShell Code": "powershell",
        "PowerShell": "powershell",
        "XML Code": "xml",
        "XML": "xml",
        "Command line code": "",
    }

    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        if (
            stripped in heading_to_lang
            and i + 1 < n
            and set(lines[i + 1].strip()) == {"-"}
        ):
            heading = stripped
            lang = heading_to_lang[heading]

            out.append(line)
            out.append(lines[i + 1])
            i += 2

            while i < n and lines[i].strip() == "":
                out.append(lines[i])
                i += 1

            code_lines = []
            while i < n and (lines[i].startswith("    ") or lines[i].startswith("\t")):
                code_lines.append(lines[i])
                i += 1

            if code_lines:
                dedented = []
                for ln in code_lines:
                    if ln.startswith("    "):
                        dedented.append(ln[4:])
                    elif ln.startswith("\t"):
                        dedented.append(ln[1:])
                    else:
                        dedented.append(ln)

                raw_code = "\n".join(dedented).strip("\n")

                if lang == "json":
                    try:
                        parsed = json.loads(raw_code)
                        raw_code = json.dumps(parsed, indent=2)
                    except Exception:
                        raw_code = pseudo_pretty_json(raw_code)

                out.append("")
                if lang:
                    out.append(f"```{lang}")
                else:
                    out.append("```")
                out.extend(raw_code.splitlines())
                out.append("```")
                out.append("")
        else:
            out.append(line)
            i += 1

    return "\n".join(out)

# --------------------------------------------------
# Main conversion function (DOCX bytes → markdown + images dict)
# --------------------------------------------------

STYLE_MAP = """
p[style-name='Note'] => div.upn-note > p:fresh
p[style-name='Important'] => div.upn-important > p:fresh
p[style-name='Action'] => div.upn-action > p:fresh
p[style-name='Code JSON'] => pre > code.language-json
"""


def convert_docx_to_markdown(docx_bytes: bytes, work_item_number: str, language: str):
    # 1) Extract metadata from first table
    doc_for_meta = Document(BytesIO(docx_bytes))
    yaml_meta = extract_metadata(doc_for_meta)
    front_matter_text = build_front_matter(yaml_meta)

    # 2) Prepare image store and converter
    image_store = {}  # path_in_zip -> bytes
    image_converter = make_image_converter(image_store, work_item_number, language)

    # 3) Build a body-only DOCX (remove first table)
    doc_body = Document(BytesIO(docx_bytes))
    if doc_body.tables:
        first_table = doc_body.tables[0]._element
        first_table.getparent().remove(first_table)

    buf = BytesIO()
    doc_body.save(buf)
    buf.seek(0)

    result = mammoth.convert_to_html(
        buf,
        convert_image=image_converter,
        style_map=STYLE_MAP,
    )
    html_body = result.value

    # 4) Apply callout divs
    html_body = apply_callout_divs(html_body)

    # 5) Convert to markdown
    body_md = convert_to_markdown(html_body).strip() + "\n"

    # 6) Post-process markdown
    body_md = fix_toc_and_anchors(body_md)
    body_md = fix_broken_images(body_md)
    body_md = fix_escaped_callouts(body_md)
    body_md = fix_code_blocks(body_md)

    final_markdown = front_matter_text + body_md
    return final_markdown, image_store

